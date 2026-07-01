# ─────────────────────────────────────────────────────────────────────────────
# PDF / DOCX Parser — Robust multi-format text extraction
# ─────────────────────────────────────────────────────────────────────────────
# Strategy:
#   1. pdfplumber (primary — best for text-layer PDFs)
#   2. PyMuPDF/fitz (fallback — handles complex layouts and some scanned PDFs)
#   3. python-docx (DOCX files)
#   4. Plain text passthrough (TXT files)
# Each extraction is accompanied by quality metrics so downstream code
# can warn when text density is suspiciously low (scanned PDF with no OCR).
# ─────────────────────────────────────────────────────────────────────────────

import logging
import re
import unicodedata
from pathlib import Path
from typing import Optional

from pydantic import BaseModel

logger = logging.getLogger(__name__)


class ParsedDocument(BaseModel):
    """Result of parsing a resume or JD file."""
    raw_text: str
    extraction_method: str       # "pdfplumber" | "pymupdf" | "docx" | "txt"
    page_count: int = 0
    char_count: int = 0
    extraction_confidence: float = 1.0   # Heuristic quality score 0-1
    warnings: list[str] = []


def parse_file(file_path: str | Path) -> ParsedDocument:
    """
    Parse any supported file format into clean text.
    Auto-selects extraction method based on file extension and fallback logic.
    
    Args:
        file_path: Path to PDF, DOCX, or TXT file
    
    Returns:
        ParsedDocument with clean text and metadata
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    elif suffix == ".txt":
        return _parse_txt(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Use PDF, DOCX, or TXT.")


def _parse_pdf(path: Path) -> ParsedDocument:
    """Parse PDF with pdfplumber, fallback to PyMuPDF."""
    warnings = []

    # ── Attempt 1: pdfplumber ────────────────────────────────────────────────
    try:
        import pdfplumber
        full_text_parts = []
        page_count = 0

        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                full_text_parts.append(text)

        raw = "\n".join(full_text_parts)
        clean = _clean_text(raw)

        confidence = _estimate_confidence(clean, page_count)
        if confidence < 0.3:
            warnings.append(
                f"Low text density ({confidence:.2f}) — file may be image-based PDF. "
                f"Consider OCR preprocessing."
            )
            # Don't fall back yet — PyMuPDF may do better
            if confidence < 0.15:
                logger.warning(f"Extremely low confidence ({confidence:.2f}) for {path.name} — trying PyMuPDF fallback")
                raise ValueError("pdfplumber: text density too low, trying fallback")

        if len(clean) < 100:
            raise ValueError("pdfplumber: extracted text too short")

        return ParsedDocument(
            raw_text=clean,
            extraction_method="pdfplumber",
            page_count=page_count,
            char_count=len(clean),
            extraction_confidence=confidence,
            warnings=warnings,
        )

    except Exception as e:
        logger.info(f"pdfplumber failed for {path.name}: {e}. Trying PyMuPDF...")

    # ── Attempt 2: PyMuPDF (fitz) ────────────────────────────────────────────
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        page_count = len(doc)
        text_parts = []

        for page in doc:
            text = page.get_text("text")
            text_parts.append(text)

        doc.close()
        raw = "\n".join(text_parts)
        clean = _clean_text(raw)
        confidence = _estimate_confidence(clean, page_count)

        if len(clean) < 50:
            warnings.append(
                "PDF appears to be image-only. Text extraction returned minimal content. "
                "OCR is required for accurate parsing."
            )
            confidence = 0.05

        return ParsedDocument(
            raw_text=clean,
            extraction_method="pymupdf",
            page_count=page_count,
            char_count=len(clean),
            extraction_confidence=confidence,
            warnings=warnings,
        )

    except Exception as e:
        logger.error(f"PyMuPDF also failed for {path.name}: {e}")
        raise RuntimeError(f"All PDF extraction methods failed for {path.name}") from e


def _parse_docx(path: Path) -> ParsedDocument:
    """Parse DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        
        # Extract paragraphs + table cells (tables often contain experience info)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        
        raw = "\n".join(parts)
        clean = _clean_text(raw)
        
        return ParsedDocument(
            raw_text=clean,
            extraction_method="docx",
            page_count=1,       # DOCX doesn't have reliable page metadata
            char_count=len(clean),
            extraction_confidence=0.95,  # DOCX is reliably structured
            warnings=[],
        )

    except Exception as e:
        logger.error(f"DOCX parsing failed for {path.name}: {e}")
        raise RuntimeError(f"DOCX parsing failed for {path.name}") from e


def _parse_txt(path: Path) -> ParsedDocument:
    """Parse plain text file."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        clean = _clean_text(text)
        return ParsedDocument(
            raw_text=clean,
            extraction_method="txt",
            page_count=1,
            char_count=len(clean),
            extraction_confidence=1.0,
            warnings=[],
        )
    except Exception as e:
        raise RuntimeError(f"TXT parsing failed for {path.name}") from e


def _clean_text(text: str) -> str:
    """
    Normalize raw extracted text:
    - Unicode normalization (NFKC: e.g., ligatures → ASCII)
    - Strip control characters (except newlines)
    - Collapse multiple whitespace within lines
    - Collapse 3+ consecutive newlines to 2
    - Strip leading/trailing whitespace
    """
    # Unicode normalization
    text = unicodedata.normalize("NFKC", text)
    
    # Remove control characters except \n and \t
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    
    # Remove bullet unicode variants → hyphen
    text = re.sub(r"[•·▪▸►◦‣⁃]", "-", text)
    
    # Normalize tabs to spaces
    text = text.replace("\t", " ")
    
    # Collapse multiple spaces within a line
    lines = []
    for line in text.split("\n"):
        lines.append(re.sub(r"  +", " ", line).strip())
    
    text = "\n".join(lines)
    
    # Collapse 3+ consecutive blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    return text.strip()


def _estimate_confidence(text: str, page_count: int) -> float:
    """
    Heuristic quality score for extracted text.
    Based on characters-per-page ratio vs expected for a resume.
    A typical resume has 2000-4000 chars/page.
    """
    if page_count == 0:
        return 0.0
    chars_per_page = len(text) / page_count
    # Sigmoid-like normalization: 0→0, 500→0.3, 2000→0.85, 4000→1.0
    if chars_per_page < 100:
        return 0.05
    elif chars_per_page < 500:
        return 0.1 + (chars_per_page / 500) * 0.25
    elif chars_per_page < 2000:
        return 0.35 + ((chars_per_page - 500) / 1500) * 0.5
    else:
        return min(0.85 + ((chars_per_page - 2000) / 4000) * 0.15, 1.0)
